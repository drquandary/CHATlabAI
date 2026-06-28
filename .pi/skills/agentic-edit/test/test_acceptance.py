#!/usr/bin/env python3
"""
Acceptance test for the agentic-edit skill.

Generates test/draft.docx, test/edits.json, runs docx_track.py, and asserts that
the output contains w:del and w:ins elements attributed to author CHATLabAI.
"""
import json
import os
import sys
import zipfile

from docx import Document
from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))
SKILL_DIR = os.path.dirname(HERE)
SCRIPT = os.path.join(SKILL_DIR, "scripts", "docx_track.py")
TEST_DIR = HERE
DRAFT = os.path.join(TEST_DIR, "draft.docx")
EDITS = os.path.join(TEST_DIR, "edits.json")
OUT = os.path.join(TEST_DIR, "draft.tracked.docx")
CHANGELOG = os.path.join(TEST_DIR, "change-log.md")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def make_draft():
    doc = Document()
    doc.add_paragraph("This proves that architects hardwired a response in viewers.")
    doc.save(DRAFT)
    print(f"Created {DRAFT}")


def make_edits():
    edits = [
        {
            "paragraph_index": 0,
            "find": "proves that architects hardwired a response in viewers",
            "replace": "is consistent with architects recruiting a response in viewers",
            "comment": "Rule 3: hedge mechanism verb",
        }
    ]
    with open(EDITS, "w") as f:
        json.dump(edits, f, indent=2)
    print(f"Created {EDITS}")


def run_script():
    import subprocess

    cmd = [
        sys.executable,
        SCRIPT,
        "--in",
        DRAFT,
        "--edits",
        EDITS,
        "--out",
        OUT,
        "--changelog",
        CHANGELOG,
        "--date",
        "2026-06-28T14:00:00-04:00",
    ]
    print("Running:", " ".join(cmd))
    r = subprocess.run(cmd, capture_output=True, text=True)
    print("STDOUT:", r.stdout)
    if r.stderr:
        print("STDERR:", r.stderr)
    assert r.returncode == 0, f"Script failed with code {r.returncode}"
    print("Script returned 0 (OK)")


def assert_revisions():
    """Parse the output docx XML and assert w:del + w:ins with author CHATLabAI exist."""
    assert os.path.exists(OUT), f"Output not found: {OUT}"
    with zipfile.ZipFile(OUT) as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)

    dels = root.findall(f".//{{{W_NS}}}del")
    inss = root.findall(f".//{{{W_NS}}}ins")
    print(f"Found {len(dels)} <w:del> elements")
    print(f"Found {len(inss)} <w:ins> elements")

    assert len(dels) >= 1, "No <w:del> elements found in output"
    assert len(inss) >= 1, "No <w:ins> elements found in output"

    del_authors = [d.get(f"{{{W_NS}}}author") for d in dels]
    ins_authors = [i.get(f"{{{W_NS}}}author") for i in inss]
    print(f"<w:del> authors: {del_authors}")
    print(f"<w:ins> authors: {ins_authors}")

    assert all(a == "CHATLabAI" for a in del_authors), (
        f"Not all <w:del> authors are CHATLabAI: {del_authors}"
    )
    assert all(a == "CHATLabAI" for a in ins_authors), (
        f"Not all <w:ins> authors are CHATLabAI: {ins_authors}"
    )
    print("All revision authors == CHATLabAI (OK)")

    # Verify the deleted text and inserted text are correct
    del_texts = []
    for d in dels:
        for dt in d.findall(f".//{{{W_NS}}}delText"):
            del_texts.append(dt.text or "")
    ins_texts = []
    for i in inss:
        for t in i.findall(f".//{{{W_NS}}}t"):
            ins_texts.append(t.text or "")
    del_joined = "".join(del_texts)
    ins_joined = "".join(ins_texts)
    print(f"Deleted text:  {del_joined!r}")
    print(f"Inserted text: {ins_joined!r}")

    assert "proves that architects hardwired a response in viewers" in del_joined, (
        f"Deleted text mismatch: {del_joined!r}"
    )
    assert (
        "is consistent with architects recruiting a response in viewers"
        in ins_joined
    ), f"Inserted text mismatch: {ins_joined!r}"
    print("Deleted + inserted text correct (OK)")

    # Verify unique w:id values across all revision marks
    ids = []
    for d in dels:
        ids.append(d.get(f"{{{W_NS}}}id"))
    for i in inss:
        ids.append(i.get(f"{{{W_NS}}}id"))
    print(f"Revision ids: {ids}")
    assert len(ids) == len(set(ids)), f"Non-unique revision ids: {ids}"
    print("All revision ids unique (OK)")


def assert_changelog():
    assert os.path.exists(CHANGELOG), f"Change log not found: {CHANGELOG}"
    with open(CHANGELOG) as f:
        content = f.read()
    print("=== change-log.md ===")
    print(content)
    assert "Rule 3" in content, "Change log does not cite Rule 3"
    assert "CHATLabAI" in content, "Change log does not mention CHATLabAI"
    print("Change log cites Rule 3 and author CHATLabAI (OK)")


def assert_accept_reject_logic():
    """
    Verify accept/reject would produce correct text by simulating the logic:
      - Accept All  => deletions removed, insertions kept
      - Reject All  => deletions kept (as normal text), insertions removed
    """
    import tempfile

    from docx import Document

    # Parse the tracked document
    doc = Document(OUT)
    # Accept: collect text ignoring w:del, keeping w:ins text
    accepted = []
    for p in doc.paragraphs:
        for r in p.runs:
            accepted.append(r.text)
    # python-docx .runs does not include w:del/w:ins by default; we check via XML
    with zipfile.ZipFile(OUT) as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)
    # Simulate Accept All: text = ins_text + non-revision text
    accept_parts = []
    for r in root.iter(f"{{{W_NS}}}r"):
        # skip runs inside w:del
        parent = r.getparent()
        if parent is not None and etree.QName(parent).localname == "del":
            continue
        for t in r.findall(f"{{{W_NS}}}t"):
            accept_parts.append(t.text or "")
    accept_text = "".join(accept_parts)
    print(f"Accept-All text: {accept_text!r}")
    assert "is consistent with architects recruiting a response in viewers" in accept_text
    assert "hardwired" not in accept_text
    print("Accept-All simulation: replacement present, deleted text gone (OK)")

    # Simulate Reject All: text = del_text + non-revision text (no ins)
    reject_parts = []
    for r in root.iter(f"{{{W_NS}}}r"):
        parent = r.getparent()
        if parent is not None and etree.QName(parent).localname == "ins":
            continue
        # delText counts as text on reject
        for t in r.findall(f"{{{W_NS}}}delText"):
            reject_parts.append(t.text or "")
        for t in r.findall(f"{{{W_NS}}}t"):
            reject_parts.append(t.text or "")
    reject_text = "".join(reject_parts)
    print(f"Reject-All text: {reject_text!r}")
    assert "proves that architects hardwired a response in viewers" in reject_text
    assert "is consistent with architects recruiting" not in reject_text
    print("Reject-All simulation: original text restored, insertion gone (OK)")


def main():
    make_draft()
    make_edits()
    run_script()
    assert_revisions()
    assert_changelog()
    assert_accept_reject_logic()
    print("\n========================================")
    print("ALL ACCEPTANCE TESTS PASSED")
    print("========================================")


if __name__ == "__main__":
    main()
